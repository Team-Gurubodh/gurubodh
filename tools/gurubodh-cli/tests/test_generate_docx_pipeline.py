import hashlib
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from gurubodh.config import load_generate_docx_job
from gurubodh.content_identity import build_content_identity
from gurubodh.docx.export import generated_title, write_chapter_docx
from gurubodh.naming import chapter_output_filename
from gurubodh.pipelines.generate_docx import run_generate_docx_job
from gurubodh.project import ProjectContext


def base_config(root_dir, language="hi-IN"):
    return {
        "schema_version": "1.0.0",
        "pipeline": "generate-docx",
        "source": {
            "backend": "local",
            "root_dir": str(root_dir),
            "subject_dir": f"123_spand_rahasya/{language}",
        },
        "destination": {
            "backend": "local",
            "root_dir": str(root_dir),
            "subject_dir": f"123_spand_rahasya/{language}",
        },
        "naming": {
            "category_code": "CAT001",
            "subject_code": "SUB123",
            "title_slug": "spand-rahasya",
            "language": language,
        },
    }


def artifact_reference(config, filename, backend="local"):
    relative = f"chapters/text_and_metadata/{filename}"
    if backend == "local":
        return {"backend": "local", "path": relative, "url": None}
    source = config["source"]
    return {
        "backend": "r2",
        "bucket": source["bucket"],
        "key": f"{source['prefix']}/{source['subject_dir']}/{relative}",
        "url": None,
    }


def metadata_for(config, chapter_number, text, backend="local", schema_version="1.4.0"):
    naming = config["naming"]
    text_name = chapter_output_filename(
        {"naming": naming | {"version": "01", "subversion": "01"}}, chapter_number, ".txt"
    )
    metadata_name = Path(text_name).with_suffix(".json").name
    return {
        "schema_version": schema_version,
        "document": {
            "category_code": naming["category_code"],
            "subject_code": naming["subject_code"],
            "title_slug": naming["title_slug"],
            "chapter_number": f"{chapter_number:03d}",
            "version": "v01.01",
            "language": naming["language"],
        },
        "files": {"text_filename": text_name, "metadata_filename": metadata_name},
        "storage": {
            "artifacts": {
                "text": artifact_reference(config, text_name, backend),
                "metadata": artifact_reference(config, metadata_name, backend),
            }
        },
        "integrity": {
            "artifacts": {
                "text": {
                    "algorithm": "sha256",
                    "encoding": "UTF-8",
                    "line_endings": "LF",
                    "scope": "artifact-bytes",
                    "value": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                }
            }
        },
        "content_identity": build_content_identity(
            naming["category_code"], naming["subject_code"], naming["language"], text
        ),
    }


def manifest_payload(config, metadata_items):
    return {
        "schema_version": 1,
        "identity_contract_version": 1,
        "subject": {
            "category_code": config["naming"]["category_code"],
            "subject_code": config["naming"]["subject_code"],
            "language": config["naming"]["language"],
        },
        "chapters": [
            {
                "generated_chapter_number": metadata["document"]["chapter_number"],
                "content_key": metadata["content_identity"]["content_key"],
                "normalized_content_sha256": metadata["content_identity"]["normalized_content_sha256"],
                "metadata_artifact": metadata["storage"]["artifacts"]["metadata"],
                "text_artifact": metadata["storage"]["artifacts"]["text"],
            }
            for metadata in metadata_items
        ],
    }


def state_payload(manifest_bytes, metadata_items, state="succeeded"):
    return {
        "schema_version": 1,
        "job_id": "test-job",
        "state": state,
        "publication": {
            "state": "succeeded" if state == "succeeded" else state,
            "canonical_manifest": {
                "sha256": hashlib.sha256(manifest_bytes).hexdigest(),
                "chapter_numbers": [item["document"]["chapter_number"] for item in metadata_items],
            },
        },
        "chapters": [
            {
                "chapter_number": item["document"]["chapter_number"],
                "state": "succeeded",
                "proofreading": {"canonical_content_key": item["content_identity"]["content_key"]},
            }
            for item in metadata_items
        ],
    }


def write_local_release(root_dir, config, chapter_texts, schema_version="1.4.0"):
    subject = Path(root_dir) / config["source"]["subject_dir"]
    artifacts = subject / "chapters" / "text_and_metadata"
    artifacts.mkdir(parents=True, exist_ok=True)
    metadata_items = []
    for number, text in enumerate(chapter_texts, start=1):
        metadata = metadata_for(config, number, text, schema_version=schema_version)
        metadata_items.append(metadata)
        (artifacts / metadata["files"]["text_filename"]).write_text(text, encoding="utf-8")
        (artifacts / metadata["files"]["metadata_filename"]).write_text(
            json.dumps(metadata, ensure_ascii=False) + "\n", encoding="utf-8"
        )
    manifest = manifest_payload(config, metadata_items)
    manifest_path = subject / "chapters" / "chapter_content_manifest.json"
    manifest_bytes = (json.dumps(manifest, ensure_ascii=False, indent=2) + "\n").encode("utf-8")
    manifest_path.write_bytes(manifest_bytes)
    state_path = subject / "run_state" / "prep-subject" / "job-state.json"
    state_path.parent.mkdir(parents=True, exist_ok=True)
    state_path.write_text(
        json.dumps(state_payload(manifest_bytes, metadata_items), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return subject, metadata_items


class FakeR2Client:
    def __init__(self, objects=None, fail_docx_upload=False):
        self.objects = dict(objects or {})
        self.fail_docx_upload = fail_docx_upload
        self.events = []

    def prefix_has_objects(self, bucket, prefix):
        return any(key.startswith(prefix) for key in self.objects)

    def exists(self, bucket, key):
        return key in self.objects

    def list_keys(self, bucket, prefix):
        return sorted(key for key in self.objects if key.startswith(prefix))

    def download_file(self, bucket, key, path):
        self.events.append(("download", key))
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        value = self.objects[key]
        Path(path).write_bytes(value if isinstance(value, bytes) else value.encode("utf-8"))

    def upload_file(self, path, bucket, key):
        self.events.append(("upload", key))
        if self.fail_docx_upload and key.endswith(".docx"):
            raise RuntimeError("simulated DOCX upload failure")
        self.objects[key] = Path(path).read_bytes()

    def delete_keys(self, bucket, keys):
        self.events.append(("delete_keys", tuple(keys)))
        for key in keys:
            self.objects.pop(key, None)
        return list(keys)

    def delete_prefix(self, bucket, prefix):
        self.events.append(("delete_prefix", prefix))
        keys = self.list_keys(bucket, prefix)
        for key in keys:
            self.objects.pop(key, None)
        return keys


def r2_config():
    config = base_config("unused")
    location = {
        "backend": "r2",
        "bucket": "gurubodh-library-dev",
        "prefix": "cms_library",
        "subject_dir": "123_spand_rahasya/hi-IN",
        "url_base": None,
    }
    config["source"] = dict(location)
    config["destination"] = dict(location)
    return config


def r2_release_objects(config, chapter_texts):
    metadata_items = [
        metadata_for(config, index, text, backend="r2")
        for index, text in enumerate(chapter_texts, start=1)
    ]
    manifest = manifest_payload(config, metadata_items)
    manifest_bytes = (json.dumps(manifest, ensure_ascii=False, indent=2) + "\n").encode("utf-8")
    root = f"cms_library/{config['source']['subject_dir']}"
    objects = {
        f"{root}/chapters/chapter_content_manifest.json": manifest_bytes,
        f"{root}/run_state/prep-subject/job-state.json": (
            json.dumps(state_payload(manifest_bytes, metadata_items), ensure_ascii=False, indent=2) + "\n"
        ).encode("utf-8"),
    }
    for metadata, text in zip(metadata_items, chapter_texts):
        objects[metadata["storage"]["artifacts"]["metadata"]["key"]] = (
            json.dumps(metadata, ensure_ascii=False) + "\n"
        ).encode("utf-8")
        objects[metadata["storage"]["artifacts"]["text"]["key"]] = text.encode("utf-8")
    return objects


class GenerateDocxPipelineTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.context = ProjectContext(
            root=Path(self.temp_dir.name), legacy_converter=Path(self.temp_dir.name) / "converter.js"
        )

    def load(self, config):
        path = Path(self.temp_dir.name) / "generate-docx.json"
        path.write_text(json.dumps(config), encoding="utf-8")
        return load_generate_docx_job(path), path

    def test_local_generation_uses_manifest_order_title_and_exact_body_mapping(self):
        config = base_config(self.temp_dir.name)
        texts = ["पहली पंक्ति।\nदूसरी पंक्ति।\n\nदूसरा अनुच्छेद।\n", "अंतिम अध्याय।\n"]
        subject, metadata_items = write_local_release(self.temp_dir.name, config, texts)
        canonical_before = {
            path: hashlib.sha256(path.read_bytes()).hexdigest()
            for path in (subject / "chapters" / "text_and_metadata").iterdir()
        }
        loaded, config_path = self.load(config)

        result = run_generate_docx_job(
            self.context, loaded, config_path=config_path, progress=lambda _: None
        )

        output = subject / "chapters" / "msword"
        expected_names = [Path(item["files"]["text_filename"]).with_suffix(".docx").name for item in metadata_items]
        self.assertEqual(sorted(path.name for path in output.glob("*.docx")), sorted(expected_names))
        first = Document(output / expected_names[0])
        self.assertEqual(first.paragraphs[0].text, "spand-rahasya: prabodhan 001")
        self.assertEqual([p.text for p in first.paragraphs[1:]], ["पहली पंक्ति।\nदूसरी पंक्ति।", "दूसरा अनुच्छेद।"])
        manifest = json.loads((output / "docx_manifest.json").read_text(encoding="utf-8"))
        self.assertEqual([item["chapter_number"] for item in manifest["chapters"]], ["001", "002"])
        self.assertEqual(manifest["counts"], {"chapter_count": 2, "docx_file_count": 2})
        self.assertEqual(manifest["title_template"]["version"], "1.0.0")
        self.assertEqual(result["processed_chapter_count"], 2)
        self.assertFalse((subject / "full_subject").exists())
        canonical_after = {
            path: hashlib.sha256(path.read_bytes()).hexdigest()
            for path in (subject / "chapters" / "text_and_metadata").iterdir()
        }
        self.assertEqual(canonical_before, canonical_after)
        reports = list((subject / "run_reports" / "generate-docx").glob("generate-docx-*.json"))
        self.assertEqual(len(reports), 1)

    def test_docx_bytes_are_deterministic_for_marathi(self):
        text = "पहिली ओळ।\nदुसरी ओळ।\n\nदुसरा परिच्छेद।\n"
        title = generated_title("spand-rahasya", "001")
        first = Path(self.temp_dir.name) / "first.docx"
        second = Path(self.temp_dir.name) / "second.docx"
        write_chapter_docx(first, text, title, "mr-IN")
        write_chapter_docx(second, text, title, "mr-IN")
        self.assertEqual(hashlib.sha256(first.read_bytes()).hexdigest(), hashlib.sha256(second.read_bytes()).hexdigest())
        self.assertEqual(Document(first).paragraphs[0].text, title)

    def test_without_overwrite_existing_output_fails_and_writes_failure_audit(self):
        config = base_config(self.temp_dir.name)
        subject, _ = write_local_release(self.temp_dir.name, config, ["अध्याय।\n"])
        output = subject / "chapters" / "msword"
        output.mkdir(parents=True)
        previous = output / "previous.docx"
        previous.write_bytes(b"previous")
        loaded, config_path = self.load(config)

        with self.assertRaisesRegex(SystemExit, "already exists"):
            run_generate_docx_job(self.context, loaded, config_path=config_path, progress=lambda _: None)

        self.assertEqual(previous.read_bytes(), b"previous")
        report = next((subject / "run_reports" / "generate-docx").glob("*.json"))
        self.assertEqual(json.loads(report.read_text(encoding="utf-8"))["run_identity"]["status"], "failed")

    def test_generation_failure_preserves_existing_set_and_unrelated_artifacts(self):
        config = base_config(self.temp_dir.name)
        subject, _ = write_local_release(self.temp_dir.name, config, ["अध्याय।\n"])
        output = subject / "chapters" / "msword"
        output.mkdir(parents=True)
        previous = output / "previous.docx"
        previous.write_bytes(b"previous")
        semantic = subject / "chapters" / "semantic_chunks" / "keep.json"
        semantic.parent.mkdir(parents=True)
        semantic.write_text("keep", encoding="utf-8")
        loaded, config_path = self.load(config)

        with patch("gurubodh.pipelines.generate_docx.write_chapter_docx", side_effect=ValueError("bad docx")):
            with self.assertRaisesRegex(ValueError, "bad docx"):
                run_generate_docx_job(
                    self.context, loaded, overwrite=True, config_path=config_path, progress=lambda _: None
                )

        self.assertEqual(previous.read_bytes(), b"previous")
        self.assertEqual(semantic.read_text(encoding="utf-8"), "keep")

    def test_overwrite_replaces_only_msword_and_accepts_legacy_metadata(self):
        config = base_config(self.temp_dir.name)
        subject, metadata_items = write_local_release(
            self.temp_dir.name, config, ["शुद्ध पाठ।\n"], schema_version="1.3.0"
        )
        metadata_path = subject / "chapters" / "text_and_metadata" / metadata_items[0]["files"]["metadata_filename"]
        legacy = json.loads(metadata_path.read_text(encoding="utf-8"))
        legacy["files"]["msword_filename"] = "old.docx"
        legacy["storage"]["artifacts"]["msword"] = {
            "backend": "local", "path": "chapters/msword/old.docx", "url": None
        }
        metadata_path.write_text(json.dumps(legacy, ensure_ascii=False) + "\n", encoding="utf-8")
        output = subject / "chapters" / "msword"
        output.mkdir(parents=True)
        (output / "old.docx").write_bytes(b"old")
        unrelated = subject / "notes.txt"
        unrelated.write_text("keep", encoding="utf-8")
        loaded, config_path = self.load(config)

        run_generate_docx_job(
            self.context, loaded, overwrite=True, config_path=config_path, progress=lambda _: None
        )

        self.assertFalse((output / "old.docx").exists())
        self.assertTrue((output / "docx_manifest.json").is_file())
        self.assertEqual(unrelated.read_text(encoding="utf-8"), "keep")

    def test_identity_mismatch_fails_before_overwrite_replacement(self):
        config = base_config(self.temp_dir.name)
        subject, metadata_items = write_local_release(self.temp_dir.name, config, ["अध्याय।\n"])
        metadata_path = subject / "chapters" / "text_and_metadata" / metadata_items[0]["files"]["metadata_filename"]
        metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
        metadata["document"]["title_slug"] = "different-title"
        metadata_path.write_text(json.dumps(metadata, ensure_ascii=False) + "\n", encoding="utf-8")
        output = subject / "chapters" / "msword"
        output.mkdir(parents=True)
        previous = output / "previous.docx"
        previous.write_bytes(b"previous")
        loaded, config_path = self.load(config)

        with self.assertRaisesRegex(SystemExit, "metadata subject identity does not match"):
            run_generate_docx_job(
                self.context, loaded, overwrite=True, config_path=config_path, progress=lambda _: None
            )

        self.assertEqual(previous.read_bytes(), b"previous")

    def test_incomplete_prep_and_revalidation_failure_preserve_existing_output(self):
        config = base_config(self.temp_dir.name)
        subject, _ = write_local_release(self.temp_dir.name, config, ["अध्याय।\n"])
        output = subject / "chapters" / "msword"
        output.mkdir(parents=True)
        previous = output / "previous.docx"
        previous.write_bytes(b"previous")
        state_path = subject / "run_state" / "prep-subject" / "job-state.json"
        state = json.loads(state_path.read_text(encoding="utf-8"))
        state["state"] = "publishing"
        state["publication"]["state"] = "publishing"
        state_path.write_text(json.dumps(state), encoding="utf-8")
        loaded, config_path = self.load(config)
        with self.assertRaisesRegex(SystemExit, "latest prep-subject job is not succeeded"):
            run_generate_docx_job(
                self.context, loaded, overwrite=True, config_path=config_path, progress=lambda _: None
            )
        self.assertEqual(previous.read_bytes(), b"previous")

        state["state"] = "succeeded"
        state["publication"]["state"] = "succeeded"
        state_path.write_text(json.dumps(state), encoding="utf-8")
        with patch(
            "gurubodh.pipelines.generate_docx.revalidate_source_release",
            side_effect=SystemExit("canonical changed"),
        ):
            with self.assertRaisesRegex(SystemExit, "canonical changed"):
                run_generate_docx_job(
                    self.context, loaded, overwrite=True, config_path=config_path, progress=lambda _: None
                )
        self.assertEqual(previous.read_bytes(), b"previous")

    def test_r2_overwrite_removes_readiness_first_and_publishes_manifest_last(self):
        config = r2_config()
        objects = r2_release_objects(config, ["पहला।\n", "दूसरा।\n"])
        root = f"cms_library/{config['source']['subject_dir']}"
        old_manifest = f"{root}/chapters/msword/docx_manifest.json"
        old_docx = f"{root}/chapters/msword/old.docx"
        semantic = f"{root}/chapters/semantic_chunks/keep.json"
        objects.update({old_manifest: b"old manifest", old_docx: b"old docx", semantic: b"keep"})
        client = FakeR2Client(objects)
        loaded, config_path = self.load(config)

        result = run_generate_docx_job(
            self.context,
            loaded,
            overwrite=True,
            config_path=config_path,
            r2_client=client,
            progress=lambda _: None,
        )

        msword_uploads = [
            event[1]
            for event in client.events
            if event[0] == "upload" and "/chapters/msword/" in event[1]
        ]
        self.assertTrue(msword_uploads[-1].endswith("docx_manifest.json"))
        delete_manifest_index = next(
            index
            for index, event in enumerate(client.events)
            if event[0] == "delete_keys" and old_manifest in event[1]
        )
        first_docx_upload_index = next(
            index
            for index, event in enumerate(client.events)
            if event[0] == "upload" and event[1].endswith(".docx")
        )
        self.assertLess(delete_manifest_index, first_docx_upload_index)
        self.assertEqual(client.objects[semantic], b"keep")
        self.assertTrue(result["publication"]["manifest_published_last"])

    def test_interrupted_r2_overwrite_has_no_readiness_manifest_and_uploads_failure_report(self):
        config = r2_config()
        objects = r2_release_objects(config, ["पहला।\n"])
        root = f"cms_library/{config['source']['subject_dir']}"
        manifest_key = f"{root}/chapters/msword/docx_manifest.json"
        objects[manifest_key] = b"old manifest"
        client = FakeR2Client(objects, fail_docx_upload=True)
        loaded, config_path = self.load(config)

        with self.assertRaisesRegex(SystemExit, "simulated DOCX upload failure"):
            run_generate_docx_job(
                self.context,
                loaded,
                overwrite=True,
                config_path=config_path,
                r2_client=client,
                progress=lambda _: None,
            )

        self.assertNotIn(manifest_key, client.objects)
        self.assertTrue(any("/run_reports/generate-docx/" in key for key in client.objects))

    def test_invalid_dedicated_job_fails_before_pipeline(self):
        config = base_config(self.temp_dir.name)
        config["schema_version"] = "9.9.9"
        with self.assertRaisesRegex(SystemExit, "schema_version must be 1.0.0"):
            self.load(config)
        config = base_config(self.temp_dir.name)
        config["source"]["subject_dir"] = "../hi-IN"
        with self.assertRaisesRegex(SystemExit, "must not contain"):
            self.load(config)
        config = base_config(self.temp_dir.name)
        config["naming"]["title_slug"] = "unsafe slug"
        with self.assertRaisesRegex(SystemExit, "invalid value"):
            self.load(config)
        config = base_config(self.temp_dir.name)
        config["destination"]["backend"] = "unsupported"
        with self.assertRaisesRegex(SystemExit, "must be local or r2"):
            self.load(config)


if __name__ == "__main__":
    unittest.main()
