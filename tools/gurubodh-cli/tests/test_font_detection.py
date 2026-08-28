import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from gurubodh.legacy.font_detection import (
    UnsupportedSourceFontError,
    source_fonts,
    validate_supported_source_fonts,
)


def _part_xml(text, font_name):
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<w:root xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p><w:r><w:rPr><w:rFonts w:ascii="{font_name}" /></w:rPr><w:t>{text}</w:t></w:r></w:p>
</w:root>'''


class SourceFontPolicyTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)

    def source_docx(self, font_name="Mangal"):
        path = self.root / "source.docx"
        document = Document()
        run = document.add_paragraph("परीक्षण").runs[0]
        run.font.name = font_name
        document.save(path)
        return path

    def test_rejects_shreelipi_before_processing(self):
        path = self.source_docx("SHREE-DEV7-0708")

        with self.assertRaisesRegex(UnsupportedSourceFontError, "ShreeLipi/Sri-Lipi conversion is disabled"):
            validate_supported_source_fonts(path)

    def test_rejects_unapproved_font_family(self):
        path = self.source_docx("Unknown Legacy Family")

        with self.assertRaisesRegex(UnsupportedSourceFontError, '"Unknown Legacy Family"'):
            validate_supported_source_fonts(path)

    def test_accepts_approved_unicode_and_aps_families(self):
        for font_name in ("Mangal", "APS-DV-Prakash"):
            with self.subTest(font_name=font_name):
                validate_supported_source_fonts(self.source_docx(font_name))

    def test_rejects_font_in_an_inherited_paragraph_style(self):
        path = self.root / "styled.docx"
        document = Document()
        style = document.styles.add_style("UnsafeLegacy", 1)
        style.font.name = "ShreeLipi"
        document.add_paragraph("परीक्षण", style="UnsafeLegacy")
        document.save(path)

        with self.assertRaisesRegex(UnsupportedSourceFontError, '"ShreeLipi"'):
            validate_supported_source_fonts(path)

    def test_rejects_font_in_document_defaults(self):
        path = self.root / "default-font.docx"
        document_xml = _part_xml("परीक्षण", "")
        document_xml = document_xml.replace('<w:rPr><w:rFonts w:ascii="" /></w:rPr>', "")
        styles_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:ascii="ShreeLipi" /></w:rPr></w:rPrDefault></w:docDefaults>
</w:styles>'''
        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as package:
            package.writestr("word/document.xml", document_xml)
            package.writestr("word/styles.xml", styles_xml)

        with self.assertRaisesRegex(UnsupportedSourceFontError, '"ShreeLipi"'):
            validate_supported_source_fonts(path)

    def test_scans_every_text_bearing_docx_part(self):
        path = self.source_docx("Mangal")
        document = Document(path)
        for section in document.sections:
            for container in (section.header, section.footer):
                run = container.paragraphs[0].add_run("परीक्षण")
                run.font.name = "Mangal"
        document.save(path)
        with zipfile.ZipFile(path, "a", compression=zipfile.ZIP_DEFLATED) as package:
            for part_name in ("word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"):
                package.writestr(part_name, _part_xml("परीक्षण", "Mangal"))

        fonts = source_fonts(path)

        self.assertTrue(
            {
                "word/document.xml",
                "word/header1.xml",
                "word/footer1.xml",
                "word/footnotes.xml",
                "word/endnotes.xml",
                "word/comments.xml",
            }.issubset({font.part_name for font in fonts})
        )
        validate_supported_source_fonts(path)
