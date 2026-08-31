import hashlib
import json
import shutil
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from gurubodh.docx.export import validate_chapter_docx
from gurubodh.legacy.font_detection import UnsupportedSourceFontError
from gurubodh.lab_proofread import LAB_HEADING_2_PARAGRAPHS, run_lab_proofread
from gurubodh.proofreading import ProofreadingError, ProofreadingSettings
from gurubodh.project import ProjectContext


class FakeProofreader:
    def __init__(self, correction="सही"):
        self.calls = []
        self.correction = correction

    def proofread(self, text, progress=None):
        self.calls.append(text)
        return {
            "corrected_text": text.replace("गलत", self.correction),
            "edits": [
                {"original": "गलत", "corrected": self.correction, "category": "spelling", "reason": "वर्तनी"}
            ] if "गलत" in text else [],
            "estimated_input_tokens": 10,
            "attempts": 1,
            "throttle_seconds": 0.0,
            "usage": {"input_tokens": 10},
        }


class LabProofreadTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)
        self.context = ProjectContext(
            root=self.root,
            legacy_converter=Path(__file__).parents[1] / "scripts" / "legacy_font_convert.js",
        )

    def source_docx(self, name="source.docx", text="यह गलत वाक्य है।"):
        path = self.root / name
        document = Document()
        document.add_paragraph(text)
        document.save(path)
        return path

    def test_unicode_run_is_confined_immutable_and_has_validated_artifacts(self):
        source = self.source_docx()
        source_before = source.read_bytes()
        proofreader = FakeProofreader()
        progress = []
        result = run_lab_proofread(
            self.context, source, "hi-IN", self.root / "lab", proofreader=proofreader, progress=progress.append
        )

        run_dir = result["run_directory"]
        self.assertTrue(run_dir.is_relative_to((self.root / "lab").resolve()))
        self.assertEqual(run_dir.parent.name, "succeeded")
        self.assertRegex(result["run_id"], r"^\d{8}-\d{6}-[a-f0-9]{6}$")
        active_dir = run_dir.parent.parent / "active" / result["run_id"]
        self.assertEqual(progress[0], f"Lab proofread run ID: {result['run_id']} (output: {active_dir})")
        self.assertEqual(progress[-1], f"Lab proofread run succeeded: {run_dir}")
        self.assertFalse(active_dir.exists())
        self.assertEqual(source.read_bytes(), source_before)
        self.assertEqual(len(proofreader.calls), 1)
        manifest = json.loads(result["manifest_path"].read_text(encoding="utf-8"))
        self.assertEqual(manifest["outcome"], "succeeded")
        self.assertEqual(manifest["run_directory"], str(run_dir))
        self.assertTrue(manifest["non_canonical"])
        self.assertEqual(manifest["source"]["sha256"], hashlib.sha256(source_before).hexdigest())
        self.assertEqual(manifest["source"]["font_encoding"], "unicode")
        self.assertEqual(manifest["locale"]["language"], "hi-IN")
        self.assertEqual(manifest["command"]["name"], "lab proofread")
        self.assertTrue((run_dir / "report" / "extracted_source.txt").is_file())
        self.assertTrue((run_dir / "output" / "source_proofread.txt").is_file())
        self.assertTrue((run_dir / "output" / "source_proofread.docx").is_file())
        self.assertTrue((run_dir / "README.md").is_file())
        self.assertTrue((run_dir / "report" / "proofreading.diff.txt").is_file())
        self.assertTrue((run_dir / "report" / "proofreading_details.json").is_file())
        corrected = (run_dir / "output" / "source_proofread.txt").read_text(encoding="utf-8")
        validate_chapter_docx(run_dir / "output" / "source_proofread.docx", corrected, "source: proofread")
        self.assertEqual(Document(run_dir / "output" / "source_proofread.docx").paragraphs[1].text, "यह सही वाक्य है।")
        self.assertIn("non-canonical", (run_dir / "README.md").read_text(encoding="utf-8"))
        self.assertEqual(manifest["operator_readme"], "README.md")
        for relative, expected_sha256 in manifest["artifact_sha256"].items():
            self.assertEqual(hashlib.sha256((run_dir / relative).read_bytes()).hexdigest(), expected_sha256)

    def test_marathi_selects_the_marathi_locale(self):
        source = self.source_docx(text="हे गलत वाक्य आहे.")
        result = run_lab_proofread(self.context, source, "mr-IN", self.root / "lab", proofreader=FakeProofreader("चुकीचे"))
        manifest = json.loads(result["manifest_path"].read_text(encoding="utf-8"))
        self.assertEqual(manifest["locale"]["language"], "mr-IN")
        self.assertEqual(manifest["locale"]["instruction_template"]["id"], "mr-IN-proofreading")

    def test_exact_configured_marathi_paragraphs_render_as_heading_2(self):
        source = self.root / "source.docx"
        document = Document()
        document.add_paragraph("प्रबोधनातील स्मरणीय मुद्दे")
        document.add_paragraph("स्वामी विश्वसंदेश")
        document.add_paragraph("ही सामान्य ओळ आहे।")
        document.save(source)
        result = run_lab_proofread(self.context, source, "mr-IN", self.root / "lab", proofreader=FakeProofreader())
        output = result["run_directory"] / "output"
        corrected = (output / "source_proofread.txt").read_text(encoding="utf-8")
        docx_path = output / "source_proofread.docx"
        validate_chapter_docx(docx_path, corrected, "source: proofread", LAB_HEADING_2_PARAGRAPHS)
        paragraphs = Document(docx_path).paragraphs[1:]
        self.assertEqual([paragraph.style.name for paragraph in paragraphs], ["Heading 2", "Heading 2", "Normal"])

    def test_legacy_font_source_uses_transient_conversion_path(self):
        source = self.source_docx(text="fdn")
        document = Document(source)
        style = document.styles.add_style("APS Legacy", 1)
        style.font.name = "APS-DV-Prakash"
        document.paragraphs[0].style = "APS Legacy"
        document.save(source)
        proofreader = FakeProofreader()

        def fake_convert(source_path, font_name, converter, output_path, text_path=None, progress=None):
            shutil.copyfile(source_path, output_path)
            return {"converter_counts": {"aps": 1}, "total_nodes": 1, "total_chars": 3}

        with patch("gurubodh.lab_proofread.convert_docx", side_effect=fake_convert) as convert:
            result = run_lab_proofread(self.context, source, "hi-IN", self.root / "lab", proofreader=proofreader)

        convert.assert_called_once()
        manifest = json.loads(result["manifest_path"].read_text(encoding="utf-8"))
        self.assertEqual(manifest["source"]["font_encoding"], "aps")
        self.assertEqual(manifest["source"]["legacy_conversion"]["converter_counts"], {"aps": 1})

    def test_unsupported_font_fails_before_conversion_or_proofreading(self):
        source = self.source_docx(text="fdn")
        document = Document(source)
        document.paragraphs[0].runs[0].font.name = "SHREE-DEV7-0708"
        document.save(source)
        proofreader = FakeProofreader()

        with (
            patch("gurubodh.lab_proofread.convert_docx") as convert,
            self.assertRaisesRegex(UnsupportedSourceFontError, "conversion is disabled"),
        ):
            run_lab_proofread(self.context, source, "hi-IN", self.root / "lab", proofreader=proofreader)

        convert.assert_not_called()
        self.assertEqual(proofreader.calls, [])

    def test_over_limit_fails_before_a_proofreading_request_and_preserves_the_source(self):
        source = self.source_docx(text="बहुत लंबा पाठ")
        before = source.read_bytes()
        proofreader = FakeProofreader()
        progress = []
        with self.assertRaisesRegex(ProofreadingError, "configured proofreading limit"):
            run_lab_proofread(
                self.context,
                source,
                "hi-IN",
                self.root / "lab",
                proofreader=proofreader,
                settings=ProofreadingSettings(max_input_characters=5),
                progress=progress.append,
            )
        self.assertEqual(proofreader.calls, [])
        self.assertEqual(source.read_bytes(), before)
        runs = self.root / "lab" / "proofread" / "runs"
        manifests = list((runs / "failed").glob("*/run_manifest.json"))
        self.assertEqual(len(manifests), 1)
        failed_dir = manifests[0].parent.resolve()
        manifest = json.loads(manifests[0].read_text(encoding="utf-8"))
        self.assertEqual(manifest["outcome"], "failed")
        self.assertEqual(manifest["run_directory"], str(failed_dir))
        self.assertEqual(list((runs / "active").iterdir()), [])
        self.assertEqual(progress[-1], f"Lab proofread run failed: {failed_dir}")

    def test_terminal_503_records_safe_capacity_diagnostics_in_the_failed_lab_report(self):
        source = self.source_docx()
        failure = ProofreadingError(
            "service_unavailable",
            "Gemini service capacity is temporarily unavailable (HTTP 503 UNAVAILABLE).",
            retryable=True,
            request_attempts=3,
            request_diagnostics={
                "attempts": [
                    {"attempt": 1, "http_status": 503, "elapsed_seconds": 1.2, "retry_delay_seconds": 30, "server_retry_hint_used": False},
                    {"attempt": 2, "http_status": 503, "elapsed_seconds": 1.3, "retry_delay_seconds": 90, "server_retry_hint_used": True},
                    {"attempt": 3, "http_status": 503, "elapsed_seconds": 1.4, "server_retry_hint_used": False},
                ],
                "terminal_retry_exhaustion_reason": "service_unavailable_retry_exhausted",
            },
        )

        class FailingProofreader:
            def proofread(self, text, progress=None):
                raise failure

        with self.assertRaisesRegex(ProofreadingError, "service capacity"):
            run_lab_proofread(self.context, source, "hi-IN", self.root / "lab", proofreader=FailingProofreader())

        manifest_path = next((self.root / "lab" / "proofread" / "runs" / "failed").glob("*/run_manifest.json"))
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        self.assertEqual(manifest["request_diagnostics"]["attempts"][0]["http_status"], 503)
        self.assertEqual(
            manifest["request_diagnostics"]["terminal_retry_exhaustion_reason"],
            "service_unavailable_retry_exhausted",
        )
        report = (manifest_path.parent / "report" / "run_report.md").read_text(encoding="utf-8")
        self.assertIn("Terminal retry reason: `service_unavailable_retry_exhausted`", report)
        self.assertNotIn("यह गलत वाक्य है।", manifest_path.read_text(encoding="utf-8"))

    def test_rejects_canonical_lab_root_and_invalid_locale(self):
        source = self.source_docx()
        with self.assertRaisesRegex(ValueError, "cms_library"):
            run_lab_proofread(self.context, source, "hi-IN", self.root / "cms_library")
        with self.assertRaisesRegex(ValueError, "Unsupported language"):
            run_lab_proofread(self.context, source, "en-IN", self.root / "lab")
