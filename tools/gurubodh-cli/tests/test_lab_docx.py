import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from gurubodh.docx.export import write_chapter_docx
from gurubodh.lab_docx import run_lab_append_docx, run_lab_assemble_docx


def has_page_break(path):
    document = Document(path)
    return bool(document.element.body.xpath('.//w:br[@w:type="page"]'))


class LabDocxTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)

    def export(self, name, title, body):
        path = self.root / name
        write_chapter_docx(path, body + "\n", title, "hi-IN")
        return path

    def paragraph_values(self, path):
        return [paragraph.text for paragraph in Document(path).paragraphs]

    def test_assemble_discovers_direct_children_in_case_insensitive_natural_order(self):
        inputs = self.root / "inputs"
        inputs.mkdir()
        self.export("inputs/chapter10.docx", "Title 10", "Body 10")
        self.export("inputs/chapter2.DOCX", "Title 2", "Body 2")
        self.export("inputs/chapter1.docx", "Title 1", "Body 1")
        self.export("inputs/~$lock.docx", "Temporary", "Ignore")
        self.export("inputs/combined.docx", "Old output", "Ignore")
        nested = inputs / "nested"
        nested.mkdir()
        self.export("inputs/nested/chapter0.docx", "Nested", "Ignore")
        first_source = inputs / "chapter1.docx"
        first_source_before = first_source.read_bytes()

        result = run_lab_assemble_docx(inputs, inputs / "combined.docx", overwrite=True)

        self.assertEqual([path.name for path in result["sources"]], ["chapter1.docx", "chapter2.DOCX", "chapter10.docx"])
        self.assertEqual(
            self.paragraph_values(inputs / "combined.docx"),
            ["Title 1", "Body 1", "", "Title 2", "Body 2", "", "Title 10", "Body 10"],
        )
        self.assertTrue(has_page_break(inputs / "combined.docx"))
        self.assertEqual(first_source.read_bytes(), first_source_before)
        source_document = Document(first_source)
        assembled_document = Document(inputs / "combined.docx")
        for style_name in ("Normal", "Title", "Heading 2"):
            self.assertEqual(
                assembled_document.styles[style_name].font.name,
                source_document.styles[style_name].font.name,
            )

    def test_assemble_rejects_empty_directories_and_overwrite_conflicts(self):
        empty = self.root / "empty"
        empty.mkdir()
        with self.assertRaisesRegex(ValueError, "no eligible DOCX"):
            run_lab_assemble_docx(empty, self.root / "combined.docx")

        source = self.export("chapter.docx", "Title", "Body")
        output = self.export("combined.docx", "Existing", "Original")
        original = output.read_bytes()
        with self.assertRaisesRegex(FileExistsError, "--overwrite"):
            run_lab_assemble_docx(self.root, output)
        self.assertEqual(output.read_bytes(), original)
        run_lab_assemble_docx(self.root, output, overwrite=True)
        self.assertEqual(self.paragraph_values(output), ["Title", "Body"])
        self.assertTrue(source.exists())

    def test_assemble_rejects_missing_directories_and_invalid_docx_inputs(self):
        with self.assertRaisesRegex(FileNotFoundError, "Input directory does not exist"):
            run_lab_assemble_docx(self.root / "missing", self.root / "combined.docx")

        inputs = self.root / "inputs"
        inputs.mkdir()
        (inputs / "invalid.docx").write_bytes(b"not a DOCX")
        with self.assertRaisesRegex(ValueError, "Input is not a valid DOCX"):
            run_lab_assemble_docx(inputs, self.root / "combined.docx")

    def test_append_preserves_source_and_uses_page_break_by_default(self):
        source = self.export("source.docx", "Source title", "Source body")
        destination = self.export("destination.docx", "Destination title", "Destination body")
        source_before = source.read_bytes()

        result = run_lab_append_docx(source, destination)

        self.assertEqual(result["source"], source.resolve())
        self.assertTrue(result["page_break"])
        self.assertEqual(source.read_bytes(), source_before)
        self.assertEqual(
            self.paragraph_values(destination),
            ["Destination title", "Destination body", "", "Source title", "Source body"],
        )
        self.assertTrue(has_page_break(destination))

    def test_append_can_suppress_the_page_break(self):
        source = self.export("source.docx", "Source title", "Source body")
        destination = self.export("destination.docx", "Destination title", "Destination body")

        run_lab_append_docx(source, destination, page_break=False)

        self.assertEqual(
            self.paragraph_values(destination),
            ["Destination title", "Destination body", "Source title", "Source body"],
        )
        self.assertFalse(has_page_break(destination))

    def test_append_uses_the_destination_template_for_source_text(self):
        source = self.export("source.docx", "Source title", "Source body")
        source_document = Document(source)
        source_document.paragraphs[1].runs[0].font.name = "Source Direct Font"
        source_document.save(source)
        source_before = source.read_bytes()
        destination = self.export("destination.docx", "Destination title", "Destination body")
        destination_document = Document(destination)
        for style_name in ("Normal", "Title", "Heading 2"):
            destination_document.styles[style_name].font.name = "Destination Template Font"
        destination_document.save(destination)

        run_lab_append_docx(source, destination)

        appended_document = Document(destination)
        self.assertEqual(source.read_bytes(), source_before)
        self.assertEqual(appended_document.paragraphs[3].style.name, "Title")
        self.assertEqual(appended_document.paragraphs[4].style.name, "Normal")
        self.assertIsNone(appended_document.paragraphs[4].runs[0].font.name)
        for style_name in ("Normal", "Title", "Heading 2"):
            self.assertEqual(appended_document.styles[style_name].font.name, "Destination Template Font")

    def test_append_rejects_equal_missing_and_invalid_paths_before_replacement(self):
        source = self.export("source.docx", "Source title", "Source body")
        destination = self.export("destination.docx", "Destination title", "Destination body")
        original = destination.read_bytes()
        invalid = self.root / "invalid.docx"
        invalid.write_bytes(b"not a DOCX")

        with self.assertRaisesRegex(ValueError, "must be different"):
            run_lab_append_docx(source, source)
        with self.assertRaisesRegex(FileNotFoundError, "Source DOCX does not exist"):
            run_lab_append_docx(self.root / "missing.docx", destination)
        with self.assertRaisesRegex(ValueError, "not a valid DOCX"):
            run_lab_append_docx(invalid, destination)
        self.assertEqual(destination.read_bytes(), original)

    def test_atomic_failure_leaves_existing_destination_unchanged(self):
        source = self.export("source.docx", "Source title", "Source body")
        destination = self.export("destination.docx", "Destination title", "Destination body")
        original = destination.read_bytes()

        with patch("gurubodh.lab_docx._save_docx", side_effect=OSError("simulated save failure")):
            with self.assertRaisesRegex(OSError, "simulated save failure"):
                run_lab_append_docx(source, destination)

        self.assertEqual(destination.read_bytes(), original)
        self.assertEqual(list(self.root.glob(".destination.docx.*.docx")), [])


if __name__ == "__main__":
    unittest.main()
