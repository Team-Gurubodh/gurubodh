import hashlib
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from gurubodh.legacy.converter import convert_texts
from gurubodh.legacy.docx_converter import convert_docx, target_devanagari_font


CLI_ROOT = Path(__file__).parents[1]
FIXTURE_PATH = Path(__file__).parent / "fixtures" / "aps_prakash_golden.json"
LEGACY_CONVERTER = CLI_ROOT / "scripts" / "legacy_font_convert.js"


class ApsGoldenMappingTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.fixture = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))

    def test_fixture_pins_the_vendored_mapping_revision(self):
        mapping = self.fixture["mapping"]
        vendor_path = CLI_ROOT / mapping["vendor_file"]

        self.assertEqual(
            hashlib.sha256(vendor_path.read_bytes()).hexdigest(),
            mapping["sha256"],
        )

    def test_fixture_covers_the_required_mapping_categories(self):
        categories = {category for case in self.fixture["cases"] for category in case["categories"]}

        self.assertTrue(self.fixture["fixture_provenance"]["source_font_family"])
        self.assertTrue(self.fixture["fixture_provenance"]["source_document_identifier"])
        self.assertTrue(self.fixture["fixture_provenance"]["initial_verification"])
        self.assertTrue(self.fixture["fixture_provenance"]["confidence_note"])
        self.assertTrue(
            {
                "matra",
                "reph",
                "half-letter",
                "conjunct",
                "anusvar",
                "chandrabindu",
                "visarga",
                "nukta",
                "digits",
                "punctuation",
                "whitespace",
                "postprocess",
            }.issubset(categories)
        )

    def test_aps_golden_cases_match_exact_unicode_output(self):
        cases = self.fixture["cases"]
        converted = convert_texts(
            [case["legacy_input"] for case in cases],
            self.fixture["converter"],
            LEGACY_CONVERTER,
        )

        for case, actual in zip(cases, converted):
            with self.subTest(case=case["id"], categories=case["categories"]):
                self.assertEqual(actual, case["expected_unicode"])


class ApsDocxConversionTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)

    def test_real_conversion_joins_split_aps_runs_and_preserves_adjacent_unicode_text(self):
        source_path = self.root / "source.docx"
        output_path = self.root / "converted.docx"
        document = Document()
        aps_style = document.styles.add_style("APS Legacy", WD_STYLE_TYPE.PARAGRAPH)
        aps_style.font.name = "APS-DV-Prakash"
        paragraph = document.add_paragraph(style=aps_style)
        paragraph.add_run("ef")
        paragraph.add_run("keâ&")
        unicode_run = paragraph.add_run(" [Unicode]")
        unicode_run.font.name = "Mangal"
        document.save(source_path)

        conversion = convert_docx(
            source_path,
            target_devanagari_font(),
            LEGACY_CONVERTER,
            output_path,
        )

        converted = Document(output_path)
        converted_paragraph = converted.paragraphs[0]
        self.assertEqual(converted_paragraph.text, "र्कि [Unicode]")
        self.assertEqual(converted_paragraph.runs[0].text, "र्कि")
        self.assertEqual(converted_paragraph.runs[1].text, "")
        self.assertEqual(converted_paragraph.runs[2].text, " [Unicode]")
        self.assertEqual(converted_paragraph.runs[2].font.name, "Mangal")
        self.assertEqual(conversion["converter_counts"], {"aps": 1})


if __name__ == "__main__":
    unittest.main()
